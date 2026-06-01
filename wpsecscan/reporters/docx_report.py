"""Item #48 — DOCX (Word) report.

When `python-docx` is installed, write a structured .docx the client can
sign-and-return as compliance acknowledgement. When it isn't, fall back
to a minimal RTF file (universally readable by Word, Google Docs, and
LibreOffice).

CLI flag: --docx writes to {stem}.docx (or {stem}.rtf on fallback).
"""
from __future__ import annotations

from pathlib import Path

from ..models import ScanReport


def write(report: ScanReport, path: Path) -> None:
    try:
        _write_docx(report, path)
    except ImportError:
        rtf_path = path.with_suffix(".rtf")
        _write_rtf(report, rtf_path)


# ---------------------------------------------------------------------------
# Real DOCX via python-docx (preferred)
# ---------------------------------------------------------------------------

def _write_docx(report: ScanReport, path: Path) -> None:
    import docx  # type: ignore[import-not-found]
    from docx.shared import Pt, RGBColor  # noqa: PLC0415

    doc = docx.Document()
    title = doc.add_heading("WPSecScan security report", level=0)
    doc.add_paragraph(f"Target: {report.target}")
    doc.add_paragraph(f"Scanned: {report.scanned_at}")
    doc.add_paragraph(f"Risk score: {report.risk_score} / 100")

    doc.add_heading("Summary", level=1)
    s = report.summary
    summary_table = doc.add_table(rows=2, cols=5)
    summary_table.style = "Light List Accent 1"
    hdr = summary_table.rows[0].cells
    cells = summary_table.rows[1].cells
    for i, sev in enumerate(("critical", "high", "medium", "low", "info")):
        hdr[i].text = sev.upper()
        cells[i].text = str(s.get(sev, 0))

    doc.add_heading("Findings", level=1)
    for r in report.results:
        if not r.findings:
            continue
        doc.add_heading(r.check_name, level=2)
        for f in r.findings:
            p = doc.add_paragraph()
            run = p.add_run(f"[{f.severity.upper()}] ")
            run.bold = True
            colour_map = {
                "critical": RGBColor(0xFF, 0x52, 0x52),
                "high":     RGBColor(0xFF, 0x8A, 0x85),
                "medium":   RGBColor(0xF0, 0xC6, 0x74),
                "low":      RGBColor(0x79, 0xC0, 0xFF),
                "info":     RGBColor(0x8B, 0x94, 0x9E),
            }
            run.font.color.rgb = colour_map.get(f.severity, RGBColor(0, 0, 0))
            p.add_run(f.title)
            if f.evidence:
                doc.add_paragraph(f.evidence[:2000], style="Intense Quote")
            if f.remediation:
                doc.add_paragraph("Remediation: " + f.remediation[:2000])

    doc.add_heading("Sign-off", level=1)
    doc.add_paragraph("Reviewed by:  _____________________________________")
    doc.add_paragraph("Title / role: _____________________________________")
    doc.add_paragraph("Date:         _____________________________________")
    doc.add_paragraph("Signature:    _____________________________________")
    doc.save(str(path))


# ---------------------------------------------------------------------------
# RTF fallback (~zero deps; opens in Word, Google Docs, LibreOffice)
# ---------------------------------------------------------------------------

def _rtf_escape(s: str) -> str:
    """Escape RTF reserved chars + replace non-ASCII with Unicode tokens."""
    out = []
    for ch in s:
        cp = ord(ch)
        if ch == "\\":
            out.append("\\\\")
        elif ch == "{":
            out.append("\\{")
        elif ch == "}":
            out.append("\\}")
        elif ch == "\n":
            out.append("\\par ")
        elif cp < 128:
            out.append(ch)
        else:
            # Unicode escape: \uN? — RTF requires a fallback ASCII char.
            # Use ? as fallback. Negative for codes >= 0x8000.
            v = cp if cp < 0x8000 else cp - 0x10000
            out.append(f"\\u{v}?")
    return "".join(out)


def _write_rtf(report: ScanReport, path: Path) -> None:
    lines: list[str] = []
    lines.append(r"{\rtf1\ansi\ansicpg65001\deff0")
    lines.append(r"{\fonttbl{\f0\fnil\fcharset0 Arial;}}")
    lines.append(r"\fs28\b WPSecScan security report\b0\par")
    lines.append(rf"\fs22 Target: {_rtf_escape(report.target)}\par")
    lines.append(rf"Scanned: {_rtf_escape(report.scanned_at)}\par")
    lines.append(rf"Risk score: {report.risk_score} / 100\par\par")

    s = report.summary
    lines.append(r"\b Summary\b0\par")
    lines.append(
        f"Critical {s.get('critical', 0)} | "
        f"High {s.get('high', 0)} | Medium {s.get('medium', 0)} | "
        f"Low {s.get('low', 0)} | Info {s.get('info', 0)}\\par\\par"
    )

    lines.append(r"\b Findings\b0\par")
    for r in report.results:
        if not r.findings:
            continue
        lines.append(rf"\b {_rtf_escape(r.check_name)}\b0\par")
        for f in r.findings:
            lines.append(rf"\b [{f.severity.upper()}]\b0  {_rtf_escape(f.title)}\par")
            if f.evidence:
                lines.append(_rtf_escape(f.evidence[:1500]) + r"\par")
            if f.remediation:
                lines.append(rf"\b Fix:\b0  {_rtf_escape(f.remediation[:1500])}\par")
            lines.append(r"\par")

    lines.append(r"\par\b Sign-off\b0\par")
    lines.append(r"Reviewed by:  _____________________________________\par")
    lines.append(r"Date:         _____________________________________\par")
    lines.append(r"Signature:    _____________________________________\par")
    lines.append("}")
    # v2.8.3 H3 — atomic temp+rename via shared helper.
    from . import _atomic_write_text
    _atomic_write_text(path, "\n".join(lines))
